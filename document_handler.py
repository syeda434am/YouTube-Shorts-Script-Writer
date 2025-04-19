import io
import base64
from datetime import datetime
from docx import Document

def create_txt_download_link(script_text, filename="script.txt"):
    """Create a download link for a text file."""
    b64 = base64.b64encode(script_text.encode()).decode()
    return f'<a href="data:text/plain;base64,{b64}" download="{filename}">Download TXT</a>'

def create_docx_download_link(script_text, filename="script.docx"):
    """Create a download link for a Word document."""
    # Create document
    doc = Document()
    
    # Add title
    doc.add_heading("YouTube Short Script", 0)
    
    # Add date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
    
    # Add script text
    doc.add_paragraph()
    doc.add_paragraph(script_text)
    
    # Generate DOCX bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    b64 = base64.b64encode(docx_bytes.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download DOCX</a>'